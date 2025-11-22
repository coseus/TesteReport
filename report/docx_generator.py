# report/docx_generator.py
"""
Simple DOCX generator to mirror the PDF content structure.
"""

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
import base64

# Import section docx builders if available
try:
    from report.sections.section_1_0_confidentiality_and_legal import build_section_docx as sec10_docx
except Exception:
    sec10_docx = None
# ... try-import others similarly (we'll call generically)

def _add_logo_to_doc(doc: Document, b64str: str):
    try:
        imgdata = base64.b64decode(b64str)
        with open("/tmp/tmp_logo.png", "wb") as f:
            f.write(imgdata)
        doc.add_picture("/tmp/tmp_logo.png", width=Inches(1.8))
    except Exception:
        pass

def generate_docx_bytes(report: dict) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Cover
    doc.add_heading("PENETRATION TEST REPORT", level=1)
    doc.add_paragraph(f"Project: {report.get('project','')}")
    doc.add_paragraph(f"Client: {report.get('client','')}")
    doc.add_paragraph(f"Tester: {report.get('tester','')}")
    doc.add_paragraph(f"Version: {report.get('version','')}")
    doc.add_paragraph(f"Date: {report.get('date','')}")

    if report.get("logo_b64"):
        _add_logo_to_doc(doc, report["logo_b64"])

    doc.add_page_break()

    # Table of contents (simple)
    doc.add_heading("Table of Contents", level=2)
    toc_items = [
        "1.0 Confidentiality and Legal",
        "2.0 Assessment Overview",
        "3.0 Finding Severity Ratings",
        "4.0 Technical Findings",
        "5.0 Additional Reports & Scans",
        "6.0 Executive Summary",
        "7.0 Vulnerability Summary",
    ]
    for t in toc_items:
        doc.add_paragraph(t)

    # Sections: we call docx builders if present, otherwise simple placeholders
    if sec10_docx:
        try:
            sec10_docx(doc, report)
        except Exception:
            doc.add_paragraph("1.0 Confidentiality and Legal (error)")
    else:
        doc.add_heading("1.0 Confidentiality and Legal", level=2)
        doc.add_paragraph("Confidentiality and legal text placeholder.")

    # Executive summary
    doc.add_heading("6.0 Executive Summary", level=2)
    doc.add_paragraph(report.get("executive_summary",""))

    # Findings list
    if report.get("findings"):
        doc.add_heading("8.0 Findings", level=2)
        for f in report.get("findings"):
            doc.add_heading(f.get("id", "8.x") + " - " + f.get("title",""), level=3)
            doc.add_paragraph(f.get("description",""))
            doc.add_paragraph("Recommendation:")
            doc.add_paragraph(f.get("recommendation",""))
            # images: insert if any
            for img_b64 in f.get("images", []):
                try:
                    imgdata = base64.b64decode(img_b64)
                    with open("/tmp/tmp_img.png", "wb") as imf:
                        imf.write(imgdata)
                    doc.add_picture("/tmp/tmp_img.png", width=Inches(5))
                except Exception:
                    pass

    # Additional reports
    if report.get("additional_reports"):
        doc.add_heading("9.0 Additional Reports & Scans", level=2)
        for r in report["additional_reports"]:
            doc.add_heading(r.get("id","9.x") + " - " + r.get("name",""), level=3)
            doc.add_paragraph(r.get("description",""))
            for f in r.get("files", []):
                doc.add_paragraph(f"File: {f.get('name')}")
    # Save to bytes
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
